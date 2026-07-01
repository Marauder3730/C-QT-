# -*- coding: utf-8 -*-
"""Generate course design report as Word document."""
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("NEED_DOCX")
    raise

doc = Document()

def set_cn_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, "黑体", size, True)

def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, "黑体", 16, True)

def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, "黑体", 14, True)

def add_body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, "宋体", 12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        set_cn_font(r, "宋体", 12)

# ===== Cover =====
add_title("面向对象程序设计")
doc.add_paragraph()
add_title("课程设计报告", 26)
doc.add_paragraph()
doc.add_paragraph()
for line in [
    "题    目：公司员工管理系统",
    "学    院：____________________",
    "专    业：____________________",
    "班    级：____________________",
    "姓    名：汪映娴（WangYixian）",
    "学    号：____________________",
    "指导教师：____________________",
    "完成日期：2026年6月",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    set_cn_font(r, "宋体", 14)

doc.add_page_break()

# ===== Abstract =====
add_h1("摘  要")
add_body(
    "本课程设计基于 Qt 5.12.9 框架，采用 C++ 面向对象方法实现了“公司员工管理系统”。"
    "系统采用三层架构：实体类层（Employee 基类及 Engineer、Leader、ChiefEngineer 派生类）、"
    "业务逻辑层（EmployeeManager）和界面层（Qt GUI）。实现了题目的四项基本功能："
    "输入、输出、修改、删除公司员工。在此基础上扩展了登录验证、JSON 文件持久化、"
    "员工证件照、领导与工程师上下级关联、统计分析图表等功能。"
    "通过继承、多态、封装等 OOP 特性，实现了不同岗位员工的差异化管理与薪资计算。"
    "测试表明系统运行稳定，界面友好，满足课程设计基本要求并具备一定扩展性。"
)
add_body("关键词：面向对象；Qt；员工管理；继承与多态；JSON 文件存储")
doc.add_page_break()

# ===== 1 =====
add_h1("一、选题说明与任务要求")
add_h2("1.1 选题依据")
add_body(
    "根据课程设计任务书，本人学号对应题目为“公司员工管理系统”（任务编号算法：学号末两位 mod 11）。"
    "该题目要求设计一个能够管理公司人员的系统，人员包括工程师、领导、主任工程师三类，"
    "并实现输入、输出、修改、删除四项基本功能。"
)
add_body(
    "课程设计强调综合运用面向对象程序设计、文件流读写及 GUI 开发。"
    "本系统选用 Qt Creator 作为开发环境，既满足 OOP 考核要求，又实现了图形化人机交互界面。"
)
add_h2("1.2 基本功能要求对照")
table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"
hdr = ["序号", "任务书要求", "本系统实现"]
for i, t in enumerate(hdr):
    table.rows[0].cells[i].text = t
rows = [
    ("1", "输入公司员工", "新增员工对话框，支持三类岗位信息录入"),
    ("2", "输出公司员工", "主界面表格展示全部员工，支持查询筛选"),
    ("3", "修改公司员工", "选中员工后修改，双击行亦可编辑"),
    ("4", "删除公司员工", "删除前弹出确认对话框，防止误删"),
]
for i, row in enumerate(rows, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

add_h2("1.3 扩展功能（评优加分项）")
add_bullet("用户登录验证（用户名 WangYixian，密码加密存储）")
add_bullet("JSON 格式数据导入/导出及自动保存")
add_bullet("员工证件照上传与显示")
add_bullet("工程师选择直属领导，领导自动显示下属列表（组织关系关联）")
add_bullet("统计分析：岗位人数饼图、平均薪资柱状图")

# ===== 2 =====
add_h1("二、系统分析与总体设计")
add_h2("2.1 需求分析")
add_body(
    "公司需要对三类员工进行统一管理。不同岗位具有不同属性与薪资计算规则："
    "工程师按参与项目数量计奖；领导按行政级别与下属人数计薪；"
    "主任工程师按负责项目数及技术津贴计薪。系统需支持数据的持久化存储，"
    "并在图形界面中完成增删改查操作。"
)
add_h2("2.2 三层架构设计")
add_body("系统采用“界面层—业务逻辑层—实体类层”三层结构，实现高内聚、低耦合：")
add_bullet("实体类层：Employee 抽象基类，Engineer、Leader、ChiefEngineer 派生类")
add_bullet("业务逻辑层：EmployeeManager 封装增删改查、文件读写、统计与关系同步")
add_bullet("界面层：MainWindow 主界面，AddEditDialog 录入对话框，LoginDialog、StatDialog 等")
add_h2("2.3 类图关系（文字描述）")
add_body(
    "Employee 为抽象基类，定义 id、name、gender、age、hireDate、photo 等公共属性，"
    "以及 calcSalary()、getType()、getDetail()、toJson() 等纯虚/虚函数。"
    "三个派生类分别重写上述函数，实现岗位差异化行为。"
    "EmployeeManager 持有 QVector<Employee*> 容器，利用多态统一管理各类员工对象。"
)

# ===== 3 OOP =====
add_h1("三、面向对象设计说明（核心）")
add_h2("3.1 封装（Encapsulation）")
add_body(
    "Employee 及其派生类的成员变量均为 private，外部通过 public 的 get/set 方法访问，"
    "保证数据安全。业务操作不直接操作界面控件，而是通过 EmployeeManager 提供的接口完成。"
)
add_h2("3.2 继承（Inheritance）")
add_body(
    "Engineer、Leader、ChiefEngineer 均 public 继承 Employee，复用公共属性与 JSON 序列化逻辑，"
    "并扩展各自专属字段（如工程师的 projects、leaderId；领导的 department、level 等）。"
)
add_h2("3.3 多态（Polymorphism）")
add_body(
    "基类定义纯虚函数 calcSalary()、getType()、getDetail()，派生类分别实现不同算法。"
    "EmployeeManager 使用 Employee* 指针存储所有员工，调用 calcSalary() 时自动绑定到具体派生类实现，"
    "体现运行时多态。fromJson() 工厂方法根据 type 字段创建对应派生类对象。"
)
add_h2("3.4 三类员工薪资算法")
table2 = doc.add_table(rows=4, cols=3)
table2.style = "Table Grid"
for i, t in enumerate(["岗位", "计算公式", "说明"]):
    table2.rows[0].cells[i].text = t
alg = [
    ("工程师", "3000 + 项目数×500", "项目数由参与项目列表自动统计"),
    ("领导", "5000 + 2000 + 下属人数×100", "下属人数由关联工程师自动统计"),
    ("主任工程师", "4000 + 负责项目数×800 + 技术津贴", "技术津贴可单独设置"),
]
for i, row in enumerate(alg, 1):
    for j, val in enumerate(row):
        table2.rows[i].cells[j].text = val

# ===== 4 =====
add_h1("四、主要功能模块设计")
add_h2("4.1 主界面 MainWindow")
add_body(
    "主界面包含菜单栏、工具栏、左侧查询与操作区、右侧员工表格、证件照与组织关系预览区。"
    "表格显示照片、工号、姓名、性别、年龄、岗位类型、岗位详情、入职日期、月薪等列。"
    "支持按姓名/工号查询、导入导出 JSON、统计分析入口。"
)
add_h2("4.2 新增/修改对话框 AddEditDialog")
add_body(
    "采用 QStackedWidget 实现岗位切换时动态显示不同表单项。"
    "工程师可填写职称、参与项目、选择直属领导；领导可查看自动关联的下属列表；"
    "支持证件照选择与薪资自动计算。"
)
add_h2("4.3 组织关系管理")
add_body(
    "工程师录入时通过 leaderId 字段关联领导工号；EmployeeManager::syncRelationships() "
    "在增删改及加载数据后自动同步领导的下属人数。"
    "getSubordinates() 按 leaderId 查询下属工程师，relationshipSummary() 生成可读的关系描述。"
)
add_h2("4.4 文件存储")
add_body(
    "使用 JSON 格式存储于 employee.json。每个员工对象包含 type 字段标识岗位，"
    "照片以 Base64 编码存入 photo 字段。程序启动时 loadFromFile() 加载，"
    "退出时 autoSave() 自动保存。"
)

# ===== 5 =====
add_h1("五、关键源代码说明")
add_body("以下列出核心源文件及职责（完整源码见附录或提交光盘/压缩包）：")
files = [
    ("employee.h / employee.cpp", "员工抽象基类，JSON 序列化与工厂方法 fromJson"),
    ("engineer.h / engineer.cpp", "工程师类，项目列表与领导关联"),
    ("leader.h / leader.cpp", "领导类，部门、级别、下属统计"),
    ("chiefengineer.h / chiefengineer.cpp", "主任工程师类"),
    ("employeemanager.h / employeemanager.cpp", "增删改查、文件读写、统计、关系同步"),
    ("mainwindow.h / mainwindow.cpp / mainwindow.ui", "主界面与交互逻辑"),
    ("addeditdialog.*", "员工信息录入与修改"),
    ("logindialog.*", "登录验证"),
    ("statdialog.*", "Qt Charts 统计图表"),
    ("main.cpp", "程序入口"),
    ("TheFirstTestForEMS.pro", "Qt 工程配置文件"),
]
for f, d in files:
    add_bullet(f"{f}：{d}")

add_h2("5.1 多态示例（calcSalary）")
add_body(
    "Engineer::calcSalary() 返回 3000.0 + m_projects.size() * 500.0；"
    "Leader::calcSalary() 返回 5000.0 + 2000.0 + m_subordinateCount * 100.0；"
    "ChiefEngineer::calcSalary() 返回 4000.0 + m_leadProjectCount * 800.0 + m_techAllowance。"
    "主界面刷新表格时调用 emp->calcSalary()，无需判断具体类型。"
)

# ===== 6 =====
add_h1("六、系统测试")
add_h2("6.1 测试环境")
add_bullet("操作系统：Windows 10/11")
add_bullet("开发工具：Qt Creator 4.x")
add_bullet("Qt 版本：5.12.9（MinGW 32-bit）")
add_bullet("模块：Qt Widgets、Qt Charts、Qt SVG")

add_h2("6.2 测试用例与结果")
tests = [
    ("登录测试", "输入正确用户名密码进入系统；错误密码提示并限制次数", "通过"),
    ("新增工程师", "填写完整信息并选择领导，保存后表格显示正确", "通过"),
    ("新增领导", "保存后修改该领导，下属列表显示已关联工程师", "通过"),
    ("修改员工", "双击表格行修改信息，数据更新并自动保存", "通过"),
    ("删除员工", "确认删除后记录从表格消失", "通过"),
    ("查询功能", "按姓名/工号关键词筛选，显示全部恢复完整列表", "通过"),
    ("导入导出", "导出 JSON 后导入，数据一致", "通过"),
    ("统计分析", "饼图、柱状图正确反映当前数据", "通过"),
    ("程序重启", "关闭再打开，employee.json 数据正确加载", "通过"),
]
table3 = doc.add_table(rows=len(tests)+1, cols=4)
table3.style = "Table Grid"
for i, t in enumerate(["测试项", "测试步骤", "预期结果", "结论"]):
    table3.rows[0].cells[i].text = t
for i, row in enumerate(tests, 1):
    for j, val in enumerate(row):
        table3.rows[i].cells[j].text = val

# ===== 7 =====
add_h1("七、总结与体会")
add_body(
    "本次课程设计完成了公司员工管理系统从需求分析、OOP 类设计、Qt 界面开发到测试的完整流程。"
    "通过 Employee 类族的设计，深入理解了封装、继承、多态的实际应用；"
    "通过 EmployeeManager 将业务与界面分离，体会到分层架构的优势；"
    "通过 JSON 文件读写掌握了 Qt 文件流与 JSON 处理。"
)
add_body(
    "在扩展功能中实现了登录、证件照、组织关系关联和统计图表，"
    "提高了系统的实用性和完整性。开发过程中也遇到信号槽连接、命名空间、"
    "界面布局等实际问题，通过查阅文档和调试逐一解决，"
    "进一步提高了 C++ 与 Qt 的综合应用能力。"
)
add_body(
    "不足与展望：可增加数据备份提醒、Excel 导出、权限分级管理等功能；"
    "界面可进一步优化响应式布局。总体而言，本系统满足任务书基本功能要求，"
    "并在 OOP 设计、文件存储和 GUI 开发方面达到课程设计预期目标。"
)

# ===== 8 =====
add_h1("八、参考文献")
refs = [
    "[1] 谭浩强. C++程序设计（第3版）[M]. 清华大学出版社.",
    "[2] Qt Company. Qt 5.12 Documentation[EB/OL]. https://doc.qt.io/qt-5/",
    "[3] 课程设计任务（指导）书. 面向对象程序设计.",
    "[4] 张海藩, 牟永敏. 软件工程导论（第6版）[M]. 清华大学出版社.",
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    set_cn_font(r, "宋体", 12)

doc.add_page_break()
add_h1("附录 A  源程序文件清单")
for f in [
    "main.cpp", "mainwindow.h", "mainwindow.cpp", "mainwindow.ui",
    "employee.h", "employee.cpp", "engineer.h", "engineer.cpp",
    "leader.h", "leader.cpp", "chiefengineer.h", "chiefengineer.cpp",
    "employeemanager.h", "employeemanager.cpp",
    "addeditdialog.h", "addeditdialog.cpp", "addeditdialog.ui",
    "logindialog.h", "logindialog.cpp", "logindialog.ui",
    "statdialog.h", "statdialog.cpp", "uistyle.h", "uistyle.cpp",
    "TheFirstTestForEMS.pro", "employee.json", "resources/resources.qrc",
]:
    add_bullet(f)

add_h2("附录 B  提交说明")
add_body(
    "课程设计一般需提交：① 课程设计报告（本文档）；② 源程序（TheFirstTestForEMS 整个工程文件夹）；"
    "③ 可执行程序或演示截图。请将学号、姓名、班级填写至封面。"
    "源码不需全部粘贴在正文，在附录列出文件清单即可；"
    "若教师要求附核心代码，建议附 employee.h、engineer.cpp、employeemanager.cpp 中多态与 JSON 相关片段。"
)

out = r"C:\Users\Wangyixian\Projects\TheFirstTestForEMS\OOP课程设计报告_公司员工管理系统.docx"
doc.save(out)
print("SAVED:" + out)
